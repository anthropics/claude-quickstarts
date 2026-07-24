#!/usr/bin/env python3
"""Generate analyst_note.docx — the Word-format document for the demo room.

The note is synthetic and clearly marked as such. It quotes adjusted EBITDA
on the Credit Agreement basis, which is deliberately non-comparable to the
figure the company reports itself. The wiki must record both numbers with
the basis each one is measured on.

Requires python-docx.
"""

from pathlib import Path

from docx import Document

HERE = Path(__file__).parent
OUT = HERE / "example_data" / "analyst_note.docx"


def main():
    doc = Document()
    doc.add_heading("Example Capital Research — Flash Note (demo)", level=1)
    doc.add_paragraph(
        "SYNTHETIC DOCUMENT — created for this demo, not a real research note."
    )
    doc.add_heading(
        "Squarespace (SQSP) — Permira's $44 looks like an opening bid, not a closing one",
        level=2,
    )
    doc.add_paragraph("May 14, 2024")
    doc.add_paragraph(
        "Permira's $44.00/share all-cash offer values SQSP at roughly $6.9B "
        "enterprise value. Founder Anthony Casalena, General Atlantic and Accel "
        "are rolling equity rather than cashing out, which we read as insiders "
        "seeing value above the deal price. With the stock having traded above "
        "$44 within the last twelve months, we would not rule out a sweetened "
        "price before closing."
    )
    doc.add_paragraph(
        "On profitability: we estimate FY2023 Consolidated EBITDA of "
        "approximately $252M on the Credit Agreement definition, which permits "
        "addbacks beyond the company's own non-GAAP measure. Squarespace's "
        "reported adjusted EBITDA for FY2023 was $235.4M. Neither figure is "
        "comparable to unlevered free cash flow. Investors comparing leverage "
        "math across sources should check which definition is in play before "
        "quoting a multiple."
    )
    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Metric", "Basis", "FY2023 value"
    r1 = table.rows[1].cells
    r1[0].text, r1[1].text, r1[2].text = (
        "Adjusted EBITDA",
        "company non-GAAP (reported)",
        "$235.4M",
    )
    r2 = table.rows[2].cells
    r2[0].text, r2[1].text, r2[2].text = (
        "Consolidated EBITDA",
        "Credit Agreement definition (est.)",
        "~$252M",
    )
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
